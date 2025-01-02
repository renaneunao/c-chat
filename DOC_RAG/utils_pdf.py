import io  # Import io for handling BytesIO
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain_core.prompts import PromptTemplate
from pydub import AudioSegment
import moviepy.config as mp_config
import os
from dotenv import load_dotenv
import tempfile
import subprocess

# noinspection PyUnresolvedReferences
from UTILS.utils import get_client_openai, gerar_contexto, atualizar_historico

# Carregar variáveis de ambiente
load_dotenv()

# Diretório dos arquivos PDF
pdf_directory = r"C:\Users\renanb3137_00\Sicoob Central Crediminas\Anexos"

client = get_client_openai()

# Caminho do ffmpeg
ffmpeg_path = r"C:\Program Files\ffmpeg\bin\ffmpeg.exe"
AudioSegment.converter = ffmpeg_path

# Configurar moviepy para usar o ffmpeg diretamente
mp_config.change_settings({"FFMPEG_BINARY": ffmpeg_path})

# Verificar se o ffmpeg está acessível
if not os.path.exists(ffmpeg_path):
    raise FileNotFoundError(f"Erro: o caminho do ffmpeg '{ffmpeg_path}' não existe.")


def transcrever_audio(uploaded_file):
    try:
        # Criar um arquivo temporário para o arquivo carregado
        with tempfile.NamedTemporaryFile(delete=False, suffix='.opus') as temp_input_file:
            print("Arquivo temporário criado:", temp_input_file.name)
            temp_input_file.write(uploaded_file.read())
            temp_input_path = temp_input_file.name
            print("Arquivo salvo em:", temp_input_path)

        # Verificar se o arquivo é do tipo .opus e converter
        if temp_input_path.endswith('.opus'):
            print("Arquivo é do tipo .opus, realizando a conversão...")
            temp_wav_path = tempfile.mktemp(suffix=".wav")
            command = [ffmpeg_path, '-i', temp_input_path, temp_wav_path]
            try:
                subprocess.run(command, check=True)
                print(f"Áudio convertido para: {temp_wav_path}")
                file_path = temp_wav_path
            except subprocess.CalledProcessError as e:
                print(f"Erro ao converter o arquivo: {e}")
                return None
        else:
            file_path = temp_input_path

        print(f"Iniciando a transcrição do arquivo: {file_path}")

        # Transcrição do áudio
        print("Iniciando a transcrição do áudio...")
        with open(file_path, "rb") as audio_file:
            transcription = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file
            )

        print("Transcrição concluída com sucesso.")
        print("Transcrição:", transcription.text)
        return transcription.text

    except Exception as e:
        print(f"Erro ao processar o arquivo: {e}")
        return None


def list_documents(directory, st):
    """Lista arquivos PDF, Word e TXT no diretório e adiciona arquivos carregados pelo uploader."""
    if "uploaded_files" not in st.session_state:
        st.session_state.uploaded_files = []

    # Lista arquivos PDF, Word e TXT do diretório
    directory_files = [f for f in os.listdir(directory) if f.endswith('.pdf') or f.endswith('.docx') or f.endswith('.txt')]

    uploaded_files = []
    # Verifica se há arquivos carregados na sessão
    if "uploaded_files" in st.session_state:
        # Adiciona os nomes dos arquivos carregados à lista de arquivos
        uploaded_files = [file.name for file in st.session_state.uploaded_files]
    else:
        uploaded_files = []

    # Combina os arquivos do diretório e os carregados, evitando duplicados
    all_files = list(set(directory_files + uploaded_files))

    return all_files, directory_files, uploaded_files

def get_document_text(document_file, st, uploaded=False):
    text = ""
    try:
        if uploaded:
            if isinstance(document_file, io.BytesIO):
                if document_file.name.endswith('.pdf'):
                    pdf_reader = PdfReader(document_file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text() or ""
                        text += page_text
                elif document_file.name.endswith('.docx'):
                    doc = Document(document_file)
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                elif document_file.name.endswith('.txt'):
                    text = document_file.getvalue().decode("utf-8")
                elif document_file.name.endswith(('.mp3', '.mp4', '.opus')):
                    transcription = transcrever_audio(document_file)
                    if transcription:
                        text += transcription
                else:
                    raise ValueError("Tipo de arquivo não suportado.")
            else:
                raise TypeError("Expected BytesIO object for uploaded file.")
        else:
            if document_file.endswith('.pdf'):
                pdf_path = os.path.join(pdf_directory, document_file)
                pdf_reader = PdfReader(pdf_path)
                for page in pdf_reader.pages:
                    page_text = page.extract_text() or ""
                    text += page_text
            elif document_file.endswith('.docx'):
                doc_path = os.path.join(pdf_directory, document_file)
                doc = Document(doc_path)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
            elif document_file.endswith('.txt'):
                with open(os.path.join(pdf_directory, document_file), "r", encoding="utf-8") as file:
                    text += file.read()
            elif document_file.name.endswith(('.mp3', '.mp4', '.opus')):
                transcription = transcrever_audio(document_file)
                if transcription:
                    text += transcription
            else:
                raise ValueError("Tipo de arquivo não suportado.")
    except Exception as e:
        st.error(f"Erro ao processar o arquivo {document_file}: {e}")
        st.write(f"Exception type: {type(e).__name__}")
        st.write(f"Exception details: {e}")
    return text

def get_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=3000, chunk_overlap=1000)
    return text_splitter.split_text(text)

def load_or_create_vector_store(text_chunks, embeddings, st):
    if "vector_store" in st.session_state:
        vector_store = st.session_state.vector_store
    else:
        vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
        st.session_state.vector_store = vector_store
    return vector_store

def get_conversational_chain(llm):
    prompt_template = """
    Como um especialista do Sicoob Credicaf, seu papel é fornecer respostas claras, concisas e precisas, 
    ajudando a esclarecer dúvidas com base em um vasto conhecimento sobre os produtos, serviços e políticas do 
    Sicoob Credicaf. Ao responder perguntas, utilize informações detalhadas e específicas do contexto fornecido, e, 
    quando necessário, informe claramente se a resposta requer informações adicionais não disponíveis no contexto. 
    Mantenha sempre um tom profissional e informativo. Sempre traga referências.

    Contexto:\n {context}\n
    Pergunta: \n{question}\n

    Resposta:
    """
    model = llm
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return load_qa_chain(model, chain_type="stuff", prompt=prompt)


def update_cache_in_session(text_chunks, st):
    if "text_cache" not in st.session_state:
        st.session_state.text_cache = {}

    for chunk in text_chunks:
        if chunk not in st.session_state.text_cache:
            st.session_state.text_cache[chunk] = True

def load_cache_from_session(st):
    if "text_cache" in st.session_state:
        return st.session_state.text_cache
    return {}

def user_input(user_question, llm, st, chosen_pdf_id, embeddings, prefix, llm_pdf_context):
    context = gerar_contexto(st, (prefix + llm_pdf_context))
    all_files, directory_files, uploaded_files = list_documents(pdf_directory, st)

    # Process all documents
    if chosen_pdf_id == "TODOS":
        all_text_chunks = []

        # Process files from directory
        for pdf_file in directory_files:
            raw_text = get_document_text(pdf_file, st, uploaded=False)
            text_chunks = get_chunks(raw_text)
            all_text_chunks.extend(text_chunks)

        # Process uploaded files
        for pdf_file_uploaded in st.session_state.uploaded_files:
            raw_text = get_document_text(pdf_file_uploaded, st, uploaded=True)
            text_chunks = get_chunks(raw_text)
            all_text_chunks.extend(text_chunks)

        # update_cache(all_text_chunks)

        # Update session cache
        update_cache_in_session(all_text_chunks, st)


        vector_store = load_or_create_vector_store(all_text_chunks, embeddings, st)

        try:
            docs = vector_store.similarity_search(user_question)
            chain = get_conversational_chain(llm)
            response = chain({"input_documents": docs, "question": user_question, "context": context},
                             return_only_outputs=False)
            resposta = response["output_text"]

            atualizar_historico(user_question, None, resposta, st)

            if resposta:
                return resposta
        except Exception as e:
            st.error(f"Ocorreu um erro ao processar sua pergunta: {e}")

    else:
        # Determine if the file is from upload or directory
        if chosen_pdf_id in uploaded_files:
            pdf_file_obj = next(file for file in st.session_state.uploaded_files if file.name == chosen_pdf_id)
            raw_text = get_document_text(pdf_file_obj, st, uploaded=True)
        elif chosen_pdf_id in directory_files:
            raw_text = get_document_text(chosen_pdf_id, st, uploaded=False)
        else:
            st.error(f"Documento selecionado não encontrado: {chosen_pdf_id}")
            return

        # Process the selected document
        text_chunks = get_chunks(raw_text)
        # update_cache(text_chunks)

        # Update session cache
        update_cache_in_session(text_chunks , st)


        vector_store = load_or_create_vector_store(text_chunks, embeddings, st)

        try:
            docs = vector_store.similarity_search(user_question)
            chain = get_conversational_chain(llm)
            response = chain({"input_documents": docs, "question": user_question, "context": context},
                             return_only_outputs=False)
            resposta = response["output_text"]
            atualizar_historico(user_question, None, resposta, st)

            if resposta:
                return resposta

        except Exception as e:
            st.error(f"Ocorreu um erro ao processar sua pergunta para o documento {chosen_pdf_id}: {e}")
